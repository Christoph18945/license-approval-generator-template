#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""The script composes a possible license approval for a lawyer,
to check on the allowance of using xbrl taxonomy packages.

The script composes a DOCX file that serves as the license approval
for new or updated xbrl taxonomies. It contains all relevant meta
information and can be submitted e.g a lawyer of a company. The script
allows generation of license approval for different providers, though only
a JSON sample concerning the European Banking Authority (EBA) is provided
at the moment.
"""

import argparse
import datetime
from docx                     import Document
from docx.enum.dml            import MSO_THEME_COLOR_INDEX
from docx.enum.text           import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants       import RELATIONSHIP_TYPE as RT
from docx.oxml.shared         import OxmlElement, qn, CT_String
from docx.oxml.text.font      import CT_RPr
from docx.oxml.text.run       import CT_R
from docx.parts.document      import DocumentPart
from docx.shared              import Inches, Pt, RGBColor
from docx.styles.style        import _ParagraphStyle
from docx.table               import _Cell
from docx.text.run            import Run
from docx.text.paragraph      import Paragraph
from docx.text.parfmt         import ParagraphFormat
import json
import os
from typing                   import Any, MutableMapping, Tuple
import xml.etree.ElementTree  as ET
from colorama                 import init
from termcolor                import colored
from Constants                import Constants

# Usage: py -3.7 gen_lic_approval.py [-family='EBA'] [-version="3.2"]
#        py -3.7 gen_lic_approval.py [-family="lei"] [-version="2022-07-02 (REC)"]

def main() -> None:
    """Entry point of program"""
    argp: argparse.ArgumentParser = argparse.ArgumentParser(description='Generate license approval file to submit it to lawyer.')
    argp.add_argument('-family', '--family', help='The taxonomy\'s family name. E.g. EBA, BBK, ...')
    argp.add_argument('-version', '--version', help='The taxonomy\'s version')
    args: argparse.Namespace = argp.parse_args()

    objConsts: Constants = Constants()

    init() # Initialize modules for colors

    taxonomy_family_name: str = args.family
    taxonomy_version: str = args.version

    if taxonomy_family_name:
        # Find requested template via family
        located_json_file: str = ""
        for i in range(len(get_all_templates())):
            if taxonomy_family_name.lower() in get_all_templates()[i]:
                located_json_file = get_all_templates()[i]

        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # <LICENSE APPROVAL DOCUMENT>
        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        
        # Customize document
        doc: Document = Document()
        section_style: _ParagraphStyle = doc.styles['Normal']
        section_style.font.name = 'Calibri (Body)'
        section_style.font.size = Pt(12)
        section = doc.sections[0] # provide access to first section

        # ------------------------------------------------------------------------------------------------------------------
        # header section:
        # ------------------------------------------------------------------------------------------------------------------
        header: section._Header = section.header
        # table contains 1 row and 2 cells
        header_table: Any = header.add_table(1, 2, Inches(12))
        cell: _Cell
        for cell in header_table.columns[1].cells:
            cell.width = Inches(1)
        # left cell wih internal refernces internal usage
        para_l_cell: _Cell = set_paragraph(header_table, 0, 0, 0)
        run_l_cell: Run = para_l_cell.add_run(objConsts.get_header_text())
        run_l_cell.font.size = Pt(11)
        # right cell displays logo pict
        para_r_cell: _Cell = set_paragraph(header_table, 0, 1, 0)
        run_r_cell = para_r_cell.add_run()

        run_r_cell.add_picture(r"img/logo.png", width=1380000, height=520000)

        set_title(doc, WD_ALIGN_PARAGRAPH.CENTER, objConsts.get_title_main_section(), True, 13)

        # ------------------------------------------------------------------------------------------------------------------
        # meta info section about the document
        # ------------------------------------------------------------------------------------------------------------------
        doc_info_section: Any = doc.add_table(rows=3, cols=3)
        # row = 0, cells = 0,1
        set_pargraph_meta_section(doc_info_section, 0, 0, 0, WD_LINE_SPACING.SINGLE, objConsts.get_sender_form(), WD_ALIGN_PARAGRAPH.LEFT)
        set_pargraph_meta_section(doc_info_section, 0, 1, 0, WD_LINE_SPACING.SINGLE, objConsts.get_submission_text_property(), WD_ALIGN_PARAGRAPH.RIGHT)
        set_pargraph_meta_section(doc_info_section, 0, 2, 0, WD_LINE_SPACING.SINGLE, objConsts.get_submission_text_name(), WD_ALIGN_PARAGRAPH.LEFT)
        # row = 1, cells = 0,1
        set_pargraph_meta_section(doc_info_section, 1, 0, 0, WD_LINE_SPACING.SINGLE, objConsts.get_submission_to(), WD_ALIGN_PARAGRAPH.LEFT)
        set_pargraph_meta_section(doc_info_section, 1, 1, 0, WD_LINE_SPACING.SINGLE, objConsts.get_appt_or_rej_text(), WD_ALIGN_PARAGRAPH.RIGHT)
        set_pargraph_meta_section(doc_info_section, 1, 2, 0, WD_LINE_SPACING.SINGLE, "", WD_ALIGN_PARAGRAPH.LEFT)
        # row = 2, cells = 0,1                                                                                                                     # american date format
        set_pargraph_meta_section(doc_info_section, 2, 0, 0, WD_LINE_SPACING.SINGLE, objConsts.get_sub_date() + str(datetime.datetime.now().strftime("%m/%d/%Y")), WD_ALIGN_PARAGRAPH.LEFT)
        set_pargraph_meta_section(doc_info_section, 2, 1, 0, WD_LINE_SPACING.SINGLE, objConsts.get_date_appr_text(), WD_ALIGN_PARAGRAPH.RIGHT)
        set_pargraph_meta_section(doc_info_section, 2, 2, 0, WD_LINE_SPACING.SINGLE, objConsts.get_date_format(), WD_ALIGN_PARAGRAPH.LEFT)
        # set width for cells
        set_meta_section_table_cell_width(doc_info_section, 0, 3.6)
        set_meta_section_table_cell_width(doc_info_section, 1, 3.0)
        set_meta_section_table_cell_width(doc_info_section, 2, 2.2)

        set_sep_line(doc, "________________________________________________________________________", False)

        # ------------------------------------------------------------------------------------------------------------------
        # main section of the document (deals with meta information about the taxonomy)
        # ------------------------------------------------------------------------------------------------------------------
        main_table: Any = doc.add_table(rows=9, cols=2)
        set_main_section_paragraph(main_table, 0, 0, objConsts.get_third_party_name_prop())

        # Name of third party software
        # ----------------------------
        col_name_h: _Cell = main_table.rows[0].cells[1]
        if "-" not in taxonomy_family_name and taxonomy_family_name != "eba":
            col_name_h.text = iterate_over_json_file(located_json_file, "swname")
        else:
            if taxonomy_family_name == "acpr-corep":
                col_name_h.text = taxonomy_family_name.replace("-"," / ").upper() + " XBRLTaxonomy"
            elif taxonomy_family_name == "us-gaap":
                col_name_h.text = "FASB " + taxonomy_version + " SEC and US GAAP Reporting Taxonomy"
                prov_taxonomy_family_name = "FASB " + taxonomy_version + " SEC and US GAAP Reporting Taxonomy"
            else:
                col_name_h.text = iterate_over_json_file(located_json_file, "swname") # respect families with topics

        # Version number or year
        # -----------------------
        set_main_section_paragraph(main_table, 1, 0, objConsts.get_version_year_prop())
        if taxonomy_family_name == "bdp":
            # two different versions for the taxonomies provided by the Bank of Portugal.
            # therefore script call : py -3.7 gen_lic_approval.py -family="bdp" -version="2.10.1 5.0.0"
            set_main_section_paragraph(main_table, 1, 1, taxonomy_version.split(" ")[0]+" bdp v"+taxonomy_version.split(" ")[1])
        else:
            set_main_section_paragraph(main_table, 1, 1, taxonomy_version)

        # Is this a version update of 
        # previously approved software? If 
        # Yes, reason for update?
        # --------------------------------
        set_main_section_paragraph(main_table, 2, 0, objConsts.get_update_prop())
        updateOfVersion: list[str] = ["Yes","No","YES","Yes, update of the ESMA ESEF Common Recommendation (CR) version"]
        if taxonomy_family_name == "dnb-dict":
            set_main_section_paragraph(main_table, 2, 1, updateOfVersion[1])
        elif taxonomy_family_name == "us-gaap" or taxonomy_family_name == "ifrs" or taxonomy_family_name == "xbrlgl":
            set_main_section_paragraph(main_table, 2, 1, updateOfVersion[2])
        elif taxonomy_family_name == "lei":
            set_main_section_paragraph(main_table, 2, 1, updateOfVersion[3])
        else:
            set_main_section_paragraph(main_table, 2, 1, updateOfVersion[0])

        # General description of software
        # -------------------------------
        set_main_section_paragraph(main_table, 3, 0, objConsts.get_softw_desc_prop())
        set_main_section_paragraph(main_table, 3, 1, iterate_over_json_file(located_json_file, "swdescription"))
        
        # Link to software homepage
        # -------------------------
        set_main_section_paragraph(main_table, 4, 0, objConsts.get_link_property_prop())
        homepage_hyperlink = set_paragraph(main_table, 4, 1, 0)
        if taxonomy_family_name == "us-gaap":
            add_hyperlink(homepage_hyperlink, iterate_over_json_file(located_json_file, "fasbhome"), "SEC and US GAAP Taxonomies")
        elif taxonomy_family_name == "bbk":
            add_hyperlink(homepage_hyperlink, iterate_over_json_file(located_json_file, "homepage"), "Reporting - Formats(XML and XBRL)")
        elif taxonomy_family_name == "boe-banking":
            add_hyperlink(homepage_hyperlink, iterate_over_json_file(located_json_file, "homepage"), "Regulatory Reporting for the Banking Sector")
        elif taxonomy_family_name == "cipc":
            add_hyperlink(homepage_hyperlink, iterate_over_json_file(located_json_file, "homepage"), "XBRL Programs")
        elif taxonomy_family_name == "dnb-ftk":
            add_hyperlink(homepage_hyperlink, iterate_over_json_file(located_json_file, "homepage"), "Pensionsfondsen")
        elif taxonomy_family_name == "sfrdp":
            add_hyperlink(homepage_hyperlink, iterate_over_json_file(located_json_file, "home"), iterate_over_json_file(located_json_file, "home"))
        else:
            add_hyperlink(homepage_hyperlink, iterate_over_json_file(located_json_file, "homepage"), iterate_over_json_file(located_json_file, "homepage"))

        # License type (e.g. MIT, BSD, GPL)
        # ---------------------------------
        set_main_section_paragraph(main_table, 5, 0, objConsts.get_license_prop())
        str_prep_lic_type = iterate_over_json_file(located_json_file, "lictype")
        if taxonomy_family_name == "dnb-biscbs" or taxonomy_family_name == "dnb-dict" or taxonomy_family_name == "dnb-ftk":
            lic_type_hyperlink = set_paragraph(main_table, 5, 1, 0)
            add_hyperlink(lic_type_hyperlink, iterate_over_json_file(located_json_file, "lictype"), "CC-BY-4.0" )
        else:
            set_main_section_paragraph(main_table, 5, 1, str_prep_lic_type)

        # Link to website showing license:
        # --------------------------------
        set_main_section_paragraph(main_table, 6, 0, objConsts.get_link_lic_prop())
        licweb_hyperlink = set_paragraph(main_table, 6, 1, 0)
        if taxonomy_family_name == "us-gaap":
            add_hyperlink(licweb_hyperlink, iterate_over_json_file(located_json_file, "licweb"), "Terms and Conditions")
        elif taxonomy_family_name == "acpr-corep" or taxonomy_family_name == "acpr-creditimmo":
            licweb_hyperlink.add_run(iterate_over_json_file(located_json_file, "licweb"))
        elif taxonomy_family_name == "bdp":
            add_hyperlink(licweb_hyperlink, iterate_over_json_file(located_json_file, "licweb"), "Disclaimer and Copyright")
        else:
            add_hyperlink(licweb_hyperlink, iterate_over_json_file(located_json_file, "licweb"), iterate_over_json_file(located_json_file, "licweb"))
        
        if taxonomy_family_name == "boe-statistics" or taxonomy_family_name == "boe-banking" or taxonomy_family_name == "boe-insurance":
            add_hyperlink(licweb_hyperlink, iterate_over_json_file(located_json_file, "licweb1"), iterate_over_json_file(located_json_file, "licweb1"))
        # elif taxonomy_family_name == "us-gaap":
        #     add_hyperlink(licweb_hyperlink, iterate_over_json_file(located_json_file, "licweb1"), "Terms and Conditions")
        
        # Products that will introduce license?
        # --------------------------------------------
        set_main_section_paragraph(main_table, 7, 0, objConsts.get_prod_prop())
        set_main_section_paragraph(main_table, 7, 1, objConsts.get_affected_products())

        # Approximate time/version?
        # -------------------------
        set_main_section_paragraph(main_table, 8, 0, objConsts.get_time_ver_prop())
        set_main_section_paragraph(main_table, 8, 1, "2024r2")

        # ------------------------------------------------------------------------------------------------------------------
        # final section of the document
        # ------------------------------------------------------------------------------------------------------------------
        doc.add_paragraph().add_run("\nADDITIONAL COMMENTS:")
        if taxonomy_family_name == "dnb-biscbs" or taxonomy_family_name == "us gaap":
            set_additional_comment(
                doc,
                WD_ALIGN_PARAGRAPH.LEFT,
                iterate_over_json_file(located_json_file, "comment"),
                10,
                82,
                82,
                82,
                True,
                False)
        elif taxonomy_family_name == "us-gaap":
            set_additional_comment(
                doc,
                WD_ALIGN_PARAGRAPH.LEFT,
                iterate_over_json_file(located_json_file, "comment"),
                8,
                82,
                82,
                82,
                True,
                False)
        elif taxonomy_family_name == "bbk":
            add_hyperlink(
                set_additional_comment(doc, WD_ALIGN_PARAGRAPH.LEFT, "", 8, 82, 82, 82, True, False),
                iterate_over_json_file(located_json_file, "comment"),
                iterate_over_json_file(located_json_file, "comment"))
        elif taxonomy_family_name == "bdp":
            add_hyperlink(
                set_additional_comment(doc, WD_ALIGN_PARAGRAPH.LEFT, "", 8, 82, 82, 82, True, False),
                "Add direct download link", "Add direct download link")
            set_additional_comment(
                doc, WD_ALIGN_PARAGRAPH.LEFT,
                iterate_over_json_file(located_json_file, "comment"),
                10,
                82,
                82,
                82,
                True,
                False)
        elif taxonomy_family_name == "cipc":
            set_additional_comment(
                doc,
                WD_ALIGN_PARAGRAPH.LEFT,
                iterate_over_json_file(located_json_file, "comment"),
                7,
                82,
                82,
                82,
                True,
                False)
            add_hyperlink(
                set_additional_comment(doc, WD_ALIGN_PARAGRAPH.LEFT, "", 8, 82, 82, 82, True, False),
                "Add direct download link",
                "Add direct download link")
        elif taxonomy_family_name == "edinet":
            add_hyperlink(
                set_additional_comment(doc, WD_ALIGN_PARAGRAPH.LEFT, "", 8, 82, 82, 82, True, False),
                iterate_over_json_file(located_json_file, "comment"),
                iterate_over_json_file(located_json_file, "comment"))
        else:
            set_additional_comment(
                doc,
                WD_ALIGN_PARAGRAPH.LEFT,
                iterate_over_json_file(located_json_file, "comment"),
                11,
                82,
                82,
                82,
                True,
                False)

        # ------------------------------------------------------------------------------------------------------------------
        # footer section
        # ------------------------------------------------------------------------------------------------------------------
        footer: Paragraph = section.footer
        set_footer(footer, 0, "Ver: 01/2022", 10)

        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # </LICENSE APPROVAL DOCUMENT>
        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++

        # ------------------------------------------------------------------------------------------------------------------
        # Compose filename and write to file
        # ------------------------------------------------------------------------------------------------------------------
        
        # Refine approval names according to updated family
        if taxonomy_family_name == "dnb-dict":
            taxonomy_family_name = "Full " + iterate_over_json_file(located_json_file, "name").split("-")[0].replace(" DICT","") + " Data Dictionary"
        elif taxonomy_family_name == "acpr-corep":
            taxonomy_family_name = iterate_over_json_file(located_json_file, "name")+"_SUBCON"

        # Compose total filename of license approval
        if taxonomy_family_name == "Full DNB Data Dictionary":
            docx_file_name = compose_docx_file_name(
                taxonomy_family_name,
                " ",
                taxonomy_version,
                " ",
                " - Third Party Software License Approval Form ",
                " ",
                "YYYYMMDD",
                ".docx"
           )
        elif "FASB " in taxonomy_family_name:
            docx_file_name = compose_docx_file_name(
                taxonomy_family_name,
                "",
                "",
                "",
                " Reporting Taxonomy - Third Party Software License Approval Form ",
                " ",
                "YYYYMMDD",
                ".docx"
           )
        elif "us-gaap" in taxonomy_family_name:
            docx_file_name = compose_docx_file_name(
                prov_taxonomy_family_name,
                "",
                "",
                "",
                "",
                "",
                " - Third Party Software License Approval Form YYYYMMDD", ".docx"
           )    
        elif "boe-insurance" in taxonomy_family_name:
            docx_file_name = compose_docx_file_name(
                iterate_over_json_file(located_json_file, "name").replace(" INSURANCE",""),
                " ",
                taxonomy_version,
                " ",
                "Insurance Taxonomy - Third Party Software License Approval Form ",
                " ",
                "YYYYMMDD",
                ".docx"
           )
        elif "LEI" in taxonomy_family_name:
            docx_file_name = compose_docx_file_name(
                taxonomy_family_name,
                " ",
                taxonomy_version.split("-")[0],
                " ",
                "(REC) Taxonomy - Third Party Software License Approval Form",
                " ",
                "YYYYMMDD",
                ".docx"
           )            
        elif "BDP" in taxonomy_family_name:
            docx_file_name = compose_docx_file_name(
                iterate_over_json_file(located_json_file, "swname").replace(" XBRL Taxonomy",""),
                " ",
                taxonomy_version.split(" ")[1],
                " ",
                " XBRL Taxonomy - Third Party Software License Approval Form",
                " ",
                "YYYYMMDD",
                ".docx"
           )
        elif "bbk" in taxonomy_family_name:
            docx_file_name = compose_docx_file_name(
                iterate_over_json_file(located_json_file, "name"),
                " ",
                taxonomy_version,
                " ",
                " German Base XBRL Taxonomy - Third Party Software License Approval Form",
                " ",
                "YYYYMMDD",
                ".docx"
           )
        elif "EDINET" in taxonomy_family_name or "SFRDP" in taxonomy_family_name or "BOE BANKING" in taxonomy_family_name or "cmf-cl-ci" in taxonomy_family_name or "ifrs" in taxonomy_family_name:
            docx_file_name = compose_docx_file_name(
                iterate_over_json_file(located_json_file, "name"),
                " ",
                taxonomy_version,
                " ",
                "XBRL Taxonomy - Third Party Software License Approval Form",
                " ",
                "YYYYMMDD",
                ".docx"
           )
        elif "Eurofiling" in taxonomy_family_name or "EDINET" in taxonomy_family_name or "cipc" in taxonomy_family_name:
            docx_file_name = compose_docx_file_name(
                iterate_over_json_file(located_json_file, "swname"),
                " ",
                taxonomy_version,
                " ",
                "XBRL Taxonomy - Third Party Software License Approval Form",
                " ",
                "YYYYMMDD",
                ".docx"
           )
        else:
            docx_file_name = compose_docx_file_name(
                iterate_over_json_file(located_json_file, "filebasename"),
                " ",
                taxonomy_version,
                " ",
                "XBRL Taxonomy - Third Party Software License Approval Form",
                " ",
                "YYYYMMDD",
                ".docx"
           )

        # write content and save file
        save_file_path = r"./YYYY-MM-DD"
        doc.save(os.path.join(save_file_path, docx_file_name))
        print(colored("\nDocument successfully generated!", 'green')+"\n"+colored("-" * 32, 'green')+"\n"+"Your generated file: "+colored(docx_file_name, 'yellow') + " can be found at './YYYY-MM-DD/'")

def add_hyperlink(paragraph: Paragraph, url: str, text: str) -> Run:
    """
    Returns an embedded hyperlink in a text string.
    The source of this code is https://github.com/python-openxml/python-docx/issues/384.
    The code has been extended with underlining blue text color.

    Keyword arguments:
    paragraph -- paragraph where text is shown
    url       -- website
    text      -- text for embedded url
    """
    part: DocumentPart = paragraph.part # get access to document.xml.rels file and new relation id value
    relation_id: str = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink: _Element = OxmlElement('w:hyperlink') # create w:hyperlink tag and add new value
    hyperlink.set(qn('r:id'), relation_id)
    hyperlink.set(qn('w:history'), '1')
    new_run: CT_R = OxmlElement('w:r')

    rPr: CT_RPr = OxmlElement('w:rPr') # create w:rPr element
    rStyle: CT_String = OxmlElement('w:rStyle') # does not add hyperlink style
    rStyle.set(qn('w:val'), 'Hyperlink')
    # join all the xml elements, add required text to the w:r element
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # create new run object and insert hyperlink
    r: Run = paragraph.add_run()
    r._r.append(hyperlink)
    # add the styling
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return r

def get_approximate_version(path_to_artifact_database: str) -> str:
    """
    Returns the latest major release version of legacy/server products as a string.
    The approximate version is the major release version of the legacy/server products.
    This version is hardcoded in the 'C:/Projects/installer/ArtifactDatabase.xml' file.

    Keyword arguments
    path_toartifact_database -- path to the 'ArtifactDatabase.xml'
    """
    try:
        tree: ET.ElementTree = ET.parse(path_to_artifact_database)
        root: ET.Element = tree.getroot()
        for elem in root.iter():
            if 'Version' in elem.tag:
                elemAttributeMap: MutableMapping[str, str] = elem.attrib.items()
                for elem_name, elem_value in elemAttributeMap:
                    if elem_name == "MajorVersionYear":
                        return elem_value
    except OSError as e:
        print("ERROR " + str(e.errno) + ": File 'C:/Projects/installer/ArtifactDatabase.xml' not found!")

def get_all_templates() -> list:
    """Return a list with all template files plus the relative path"""
    allTemplates: list = []
    for root, directories, files in os.walk(r"./templates", topdown=False):
        for name in files:
            allTemplates.append(os.path.join(root, name))
    return allTemplates

def iterate_over_json_file(json_file: str, elem_name: str) -> str:
    """
    Return requested element out of JSON file. The data are retrieved
    from selected template of get_all_templates()

    Keyword arguments:
    json_file -- path to the json file
    elem_name -- name of the element to retrieve value
    """
    with open(json_file, "r") as data_file:
        data: dict = json.load(data_file)
        elemName: str
        elemValue: str
        for elemName, elemValue in data.items():
            if elemName == elem_name:
                return elemValue

def set_paragraph(header_table: Any, row_num: int, cell_num: int, para_num: int) -> Paragraph:
    """Return a paragraph in a table cell
    
    Keyword arguments:
    header_table -- table of header secion 
    row_num      -- row in table
    cell_num     -- cell in row
    para_num     -- set paragraph in cell
    """
    table_cell: _Cell = header_table.rows[row_num].cells[cell_num]
    para_table_cell: Paragraph = table_cell.paragraphs[para_num]
    return para_table_cell

def set_pargraph_meta_section(doc_info_table: Any, row_num: int, cell_num: int, para_num: int, format, text: str, alignment: Any) -> _Cell:
    """
    Return a pargraph for the meta inforation section.
    This section is the very top of the license approval document.

    Keyword arguments:
    doc_info_table -- table of header secion
    row_num        -- row in table
    cell_num       -- cell in row
    para_num       -- set paragraph in cell
    """
    cell: _Cell = doc_info_table.rows[row_num].cells[cell_num]
    cell_para: Paragraph = cell.paragraphs[para_num]
    cell_para_format: ParagraphFormat = cell_para.paragraph_format
    cell_para_format.line_spacing_rule = format
    cell_para.text = text
    cell_para.alignment = alignment
    return cell

def set_title(doc: Document, format: ParagraphFormat, text: str, boldness: bool, font_size: int) -> Paragraph:
    """Return paragraph with main title contained.
    
    Keyword arguments:
    doc       -- base class
    format    -- set the format of the pargraph    
    text      -- text in pargraph
    boldness  -- set boldness of text
    font_size -- set font size for the text
    """
    title_main_obj: Paragraph = doc.add_paragraph()
    title_main_obj.paragraph_format.alignment = format
    run_main_title: Run = title_main_obj.add_run(text)
    run_main_title.bold = boldness
    run_main_title.font.size = Pt(font_size)
    return title_main_obj

def set_meta_section_table_cell_width(doc_info_section: Any, colum_num: int, inche_num: float):
    """Returns width of a table cell in meta section.

    Keyword arguments:
    doc_info_section -- section with meta informatio
    column_int       -- coumn number in table
    inche_num        -- column width
    """
    cell: _Cell
    all_cells_info_sec: Tuple[_Cell, _Cell, _Cell] = doc_info_section.columns[colum_num].cells 
    for cell in all_cells_info_sec:
        cell.width = Inches(inche_num)
        return cell

def set_sep_line(doc: object, line: str, boldness: bool):
    """
    Returns the separation line. The line crosses the whole
    document vertically.

    Keyword arguments:
    doc      -- document object
    line     -- line in the document
    boldness -- set text bold
    """
    separation_line: _Cell = line
    sep_line_obj: _Cell = doc.add_paragraph().add_run(separation_line)
    sep_line_obj.bold = False
    return sep_line_obj

def set_main_section_paragraph(main_table, row_num: int, cell_num: int, text: str) -> Paragraph:
    """Returns one paragraph for the main section"""
    para: Paragraph = main_table.rows[row_num].cells[cell_num]
    para.text = text
    return para

def set_footer(footer, row_num: int, text: str, font_size: int) -> Paragraph:
    """Returns footer with text and styling"""
    footer_para: Paragraph = footer.paragraphs[row_num].add_run(text)
    footer_para.font.size = Pt(font_size)
    return footer_para

def set_additional_comment(
    doc: Document,
    alignment,
    comment: str,
    font_size: int,
    rgb_color_red: int,
    rgb_color_yellow: int,
    rgb_color_green: int,
    italic_value: bool,
    bold_value: bool) -> Paragraph:
    """Returns footer with text and styling"""
    additional_comment_obj: Paragraph = doc.add_paragraph()
    additional_comment_obj.paragraph_format.alignment = alignment
    run_main_title: Run = additional_comment_obj.add_run(comment)
    run_main_title.font.size = Pt(font_size)
    run_main_title.font.color.rgb = RGBColor(rgb_color_red, rgb_color_yellow, rgb_color_green)
    run_main_title.italic = italic_value
    run_main_title.bold = bold_value
    return additional_comment_obj

def compose_docx_file_name(taxonomy_family_name: str, ws_1: str, version, ws_2: str, general_clause: str, ws_3, ph_date: str, file_extension: str ) -> str: 
    """Compose final name for the license approval daocument
    
    taxonomy_family_name: name of the taxonomy family. E.g.: "EBA
    ws_1:                 whitespace
    version:              version of the taxonomy                
    """
    taxonomy_family_name = "EBA"
    composed_file_name: str = taxonomy_family_name + ws_1 + version + ws_2 + general_clause + ws_3 + ph_date + file_extension
    return composed_file_name

if __name__ == "__main__":
    main()